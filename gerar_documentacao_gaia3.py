#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading_with_color(doc, text, level, color):
    """Adiciona um heading com cor personalizada"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
    return heading

def add_colored_paragraph(doc, text, color=None, bold=False, size=11):
    """Adiciona parágrafo com formatação"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        if color:
            run.font.color.rgb = color
        if bold:
            run.font.bold = True
        run.font.size = Pt(size)
    return p

def shade_cell(cell, color):
    """Colore uma célula da tabela"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ===== CAPA =====
title = doc.add_heading('GAIA 3.0', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(59, 130, 246)  # Azul
    run.font.size = Pt(48)

subtitle = doc.add_paragraph('Plataforma Profissional de Marketing Digital Automatizado')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(139, 92, 246)  # Roxo

doc.add_paragraph()
doc.add_paragraph()

# Informações da capa
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = f'Guia Completo para Windows 11 Pro\nVersão 3.0\nData: {datetime.date.today().strftime("%d/%m/%Y")}'
for line in info_text.split('\n'):
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ===== ÍNDICE =====
add_heading_with_color(doc, '📋 ÍNDICE', 1, RGBColor(59, 130, 246))
toc_items = [
    '1. Introdução',
    '2. Novidades da Versão 3.0',
    '3. Pré-requisitos do Sistema',
    '4. Instalação Passo a Passo',
    '5. Login de Desenvolvedor',
    '6. Painel Administrativo',
    '7. Controle Avançado de Campanhas',
    '8. Impulsionamento de Mídias',
    '9. Guia Interativo',
    '10. Fluxo Semanal de Trabalho',
    '11. Troubleshooting',
    '12. Suporte e Comunidade',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== INTRODUÇÃO =====
add_heading_with_color(doc, '🚀 1. Introdução', 1, RGBColor(59, 130, 246))
doc.add_paragraph(
    'Bem-vindo ao GAIA 3.0, a plataforma mais avançada de marketing digital automatizado. '
    'Este guia foi criado especialmente para usuários Windows 11 Pro e iniciantes em TI. '
    'Você aprenderá tudo que precisa para dominar a plataforma em poucas horas.'
)

# ===== NOVIDADES =====
add_heading_with_color(doc, '✨ 2. Novidades da Versão 3.0', 1, RGBColor(139, 92, 246))

novidades = [
    ('🔐 Login de Desenvolvedor', 'Acesso exclusivo com credenciais seguras e autenticação de dois fatores'),
    ('👨‍💼 Painel Administrativo', 'Controle total da plataforma como desenvolvedor/criador'),
    ('📊 Controle Avançado', 'Métricas refinadas e apuradas com recomendações automáticas'),
    ('🚀 Impulsionamento de Mídias', 'Boost automático em 6 plataformas sociais'),
    ('📚 Guia Interativo', 'Aprenda dentro da plataforma com checklist e fluxo semanal'),
    ('🤖 IA Avançada', 'Geração de conteúdo com OpenAI e Google Gemini'),
    ('💬 WhatsApp Bot', 'Atendimento automático 24/7 com PLN'),
    ('💳 Pagamento Real', 'Sistema PagBrasil com PIX e Boleto'),
]

for title, desc in novidades:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(59, 130, 246)
    p.add_run(f': {desc}')

doc.add_page_break()

# ===== PRÉ-REQUISITOS =====
add_heading_with_color(doc, '⚙️ 3. Pré-requisitos do Sistema', 1, RGBColor(59, 130, 246))

doc.add_paragraph('Para usar GAIA 3.0 no Windows 11 Pro, você precisa de:')

# Tabela de pré-requisitos
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'

headers = table.rows[0].cells
headers[0].text = 'Componente'
headers[1].text = 'Versão Mínima'

shade_cell(headers[0], '3B82F6')
shade_cell(headers[1], '3B82F6')

requirements = [
    ('Windows 11 Pro', '22H2 ou superior'),
    ('Node.js', '18.0.0 ou superior'),
    ('npm ou pnpm', '10.0.0 ou superior'),
    ('Git', '2.40.0 ou superior'),
    ('MySQL', '8.0 ou superior'),
    ('RAM Disponível', '4GB mínimo (8GB recomendado)'),
]

for idx, (component, version) in enumerate(requirements, 1):
    row = table.rows[idx].cells
    row[0].text = component
    row[1].text = version

doc.add_page_break()

# ===== INSTALAÇÃO =====
add_heading_with_color(doc, '📥 4. Instalação Passo a Passo', 1, RGBColor(59, 130, 246))

steps = [
    ('Passo 1: Baixar Node.js', [
        'Acesse https://nodejs.org/',
        'Clique em "Download" (versão LTS)',
        'Execute o instalador',
        'Siga as instruções na tela',
        'Reinicie o computador',
    ]),
    ('Passo 2: Instalar Git', [
        'Acesse https://git-scm.com/',
        'Clique em "Download for Windows"',
        'Execute o instalador',
        'Use as configurações padrão',
        'Clique em "Finish"',
    ]),
    ('Passo 3: Instalar MySQL', [
        'Acesse https://dev.mysql.com/downloads/mysql/',
        'Baixe a versão 8.0 ou superior',
        'Execute o instalador MSI',
        'Configure a porta (padrão: 3306)',
        'Defina uma senha segura',
    ]),
    ('Passo 4: Clonar Repositório', [
        'Abra o Prompt de Comando (Windows + R, digite "cmd")',
        'Digite: git clone https://github.com/seu-usuario/gaia-3.0.git',
        'Aguarde o download',
        'Digite: cd gaia-3.0',
    ]),
    ('Passo 5: Instalar Dependências', [
        'No Prompt de Comando, digite: npm install -g pnpm',
        'Depois: pnpm install',
        'Aguarde a instalação (pode levar 5-10 minutos)',
    ]),
]

for step_title, step_items in steps:
    add_heading_with_color(doc, step_title, 2, RGBColor(139, 92, 246))
    for item in step_items:
        doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ===== LOGIN DE DESENVOLVEDOR =====
add_heading_with_color(doc, '🔐 5. Login de Desenvolvedor', 1, RGBColor(59, 130, 246))

doc.add_paragraph(
    'Como criador e desenvolvedor do GAIA 3.0, você tem acesso exclusivo ao painel administrativo. '
    'Aqui estão suas credenciais de acesso:'
)

# Credenciais
cred_table = doc.add_table(rows=3, cols=2)
cred_table.style = 'Light Grid Accent 1'

cred_headers = cred_table.rows[0].cells
cred_headers[0].text = 'Campo'
cred_headers[1].text = 'Valor'
shade_cell(cred_headers[0], '3B82F6')
shade_cell(cred_headers[1], '3B82F6')

cred_rows = [
    ('Usuário', 'keday49c'),
    ('Senha', 'Gaia@2024#Dev!Secure'),
]

for idx, (field, value) in enumerate(cred_rows, 1):
    row = cred_table.rows[idx].cells
    row[0].text = field
    row[1].text = value

doc.add_paragraph()
doc.add_paragraph('⚠️ IMPORTANTE: Guarde essas credenciais em local seguro. Você pode alterá-las a qualquer momento no painel.')

# Como fazer login
add_heading_with_color(doc, 'Como Fazer Login', 2, RGBColor(139, 92, 246))
login_steps = [
    'Acesse http://localhost:5173/developer-login',
    'Insira seu usuário: keday49c',
    'Insira sua senha: Gaia@2024#Dev!Secure',
    'Clique em "Acessar Painel"',
    'Você será redirecionado para o Painel Administrativo',
]
for step in login_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# ===== PAINEL ADMINISTRATIVO =====
add_heading_with_color(doc, '👨‍💼 6. Painel Administrativo', 1, RGBColor(59, 130, 246))

doc.add_paragraph(
    'O Painel Administrativo oferece controle total sobre a plataforma GAIA 3.0. '
    'Aqui você pode gerenciar tudo como desenvolvedor e criador.'
)

admin_features = [
    ('Gerenciar Usuários', 'Criar, editar, deletar e controlar permissões de usuários'),
    ('Configurações de Segurança', 'Autenticação de dois fatores, logs de auditoria, backup'),
    ('Gerenciar Credenciais', 'Adicionar, editar e remover chaves de API'),
    ('Monitorar Performance', 'Ver estatísticas de uso, erros e performance'),
    ('Gerenciar Assinaturas', 'Controlar planos, pagamentos e renovações'),
    ('Logs de Auditoria', 'Rastrear todas as ações dos usuários'),
]

for feature, desc in admin_features:
    p = doc.add_paragraph()
    run = p.add_run(feature)
    run.bold = True
    p.add_run(f': {desc}')

doc.add_page_break()

# ===== CONTROLE AVANÇADO =====
add_heading_with_color(doc, '📊 7. Controle Avançado de Campanhas', 1, RGBColor(59, 130, 246))

doc.add_paragraph(
    'O Controle Avançado oferece métricas refinadas e apuradas para cada campanha. '
    'Você pode acompanhar performance em tempo real e otimizar automaticamente.'
)

metrics = [
    ('ROAS', 'Retorno sobre Gastos em Publicidade'),
    ('ROI', 'Retorno sobre Investimento em Percentual'),
    ('CTR', 'Taxa de Cliques (Cliques / Impressões)'),
    ('CPC', 'Custo por Clique'),
    ('CPA', 'Custo por Aquisição/Conversão'),
    ('Quality Score', 'Pontuação de Qualidade (0-10)'),
    ('CPM', 'Custo por Mil Impressões'),
    ('Taxa de Conversão', 'Conversões / Cliques em Percentual'),
]

for metric, desc in metrics:
    p = doc.add_paragraph()
    run = p.add_run(metric)
    run.bold = True
    run.font.color.rgb = RGBColor(59, 130, 246)
    p.add_run(f': {desc}')

doc.add_page_break()

# ===== IMPULSIONAMENTO =====
add_heading_with_color(doc, '🚀 8. Impulsionamento de Mídias', 1, RGBColor(59, 130, 246))

doc.add_paragraph(
    'Impulsione suas postagens em 6 plataformas sociais simultaneamente. '
    'O GAIA 3.0 distribui automaticamente seu orçamento para máximo ROI.'
)

platforms = [
    'Instagram',
    'Facebook',
    'TikTok',
    'Twitter/X',
    'LinkedIn',
    'YouTube',
]

doc.add_paragraph('Plataformas Suportadas:')
for platform in platforms:
    doc.add_paragraph(platform, style='List Bullet')

doc.add_page_break()

# ===== FLUXO SEMANAL =====
add_heading_with_color(doc, '📅 10. Fluxo Semanal de Trabalho', 1, RGBColor(59, 130, 246))

doc.add_paragraph(
    'Siga este fluxo semanal para máxima eficiência e resultados consistentes.'
)

workflow_days = [
    ('SEGUNDA', 'Criar Campanha', [
        'Crie nova campanha',
        'Configure Google Ads, Meta Ads, TikTok',
        'Gere conteúdo com IA',
        'Publique automaticamente',
    ]),
    ('TERÇA A DOMINGO', 'Monitorar e Otimizar', [
        'GAIA coleta dados em tempo real',
        'Calcula métricas (CTR, CPA, ROAS)',
        'Analisa performance',
        'Gera recomendações',
        'Bot responde clientes no WhatsApp',
        'Você acompanha no Dashboard',
    ]),
    ('PRÓXIMA SEGUNDA', 'Analisar e Planejar', [
        'Veja relatório completo',
        'Analise o que funcionou',
        'Otimize a campanha',
        'Crie nova campanha',
        'Ciclo continua...',
    ]),
]

for day, title, tasks in workflow_days:
    add_heading_with_color(doc, f'{day} - {title}', 2, RGBColor(139, 92, 246))
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')

doc.add_page_break()

# ===== TROUBLESHOOTING =====
add_heading_with_color(doc, '🔧 11. Troubleshooting', 1, RGBColor(59, 130, 246))

problems = [
    ('Erro: "Port 3000 already in use"', [
        'Abra o Prompt de Comando',
        'Digite: netstat -ano | findstr :3000',
        'Copie o PID (número)',
        'Digite: taskkill /PID <número> /F',
    ]),
    ('Erro: "Cannot find module"', [
        'Abra o Prompt de Comando',
        'Digite: cd caminho/do/projeto',
        'Digite: rm -r node_modules',
        'Digite: pnpm install',
    ]),
    ('Erro: "Database connection failed"', [
        'Verifique se MySQL está rodando',
        'Abra Services (Windows + R, services.msc)',
        'Procure por "MySQL80"',
        'Clique com botão direito e selecione "Start"',
    ]),
]

for problem, solutions in problems:
    add_heading_with_color(doc, problem, 2, RGBColor(139, 92, 246))
    for solution in solutions:
        doc.add_paragraph(solution, style='List Number')

doc.add_page_break()

# ===== CONCLUSÃO =====
add_heading_with_color(doc, '🎉 Conclusão', 1, RGBColor(59, 130, 246))

doc.add_paragraph(
    'Parabéns! Você agora tem GAIA 3.0 instalado e configurado. '
    'Comece a criar campanhas, impulsionar mídias e acompanhar resultados em tempo real. '
    'Lembre-se de seguir o fluxo semanal recomendado para máxima eficiência.'
)

doc.add_paragraph()
doc.add_paragraph('Sucesso em suas campanhas de marketing! 🚀')

# Salvar documento
doc.save('/home/ubuntu/apogeu/GUIA_COMPLETO_GAIA_3.0.docx')
print('✅ Documento Word criado com sucesso: GUIA_COMPLETO_GAIA_3.0.docx')

