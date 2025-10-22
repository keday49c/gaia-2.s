#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Criar documento
doc = Document()

# Configurar estilos
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ============ CAPA ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('🚀 APOGEU\n')
title_run.font.size = Pt(48)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 102, 204)  # Azul

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Seu Pico de Sucesso em Marketing Digital\n\n')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(102, 51, 153)  # Roxo

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline_run = tagline.add_run('Guia ULTRA DETALHADO de Instalação e Configuração\n')
tagline_run.font.size = Pt(16)
tagline_run.font.italic = True

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
version_run = version.add_run('Para Windows 10 Pro - Iniciantes em TI\n\n')
version_run.font.size = Pt(14)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date.add_run('Versão 1.0 - Outubro 2024\n')
date_run.font.size = Pt(12)
date_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ============ ÍNDICE ============
doc.add_heading('📋 ÍNDICE', level=1)
indice = [
    '1. Introdução',
    '2. Pré-requisitos do Sistema',
    '3. Instalar Node.js',
    '4. Instalar Git',
    '5. Instalar MySQL',
    '6. Instalar pnpm',
    '7. Baixar Código do APOGEU',
    '8. Instalar Dependências',
    '9. Configurar Banco de Dados',
    '10. Configurar Variáveis de Ambiente',
    '11. Executar Migrações',
    '12. Iniciar Backend',
    '13. Iniciar Frontend',
    '14. Acessar Aplicação',
    '15. Executar Testes',
    '16. Configurar APIs Externas',
    '17. Build para Produção',
    '18. Criar App Desktop',
    '19. Solucionar Problemas',
    '20. Próximas Etapas',
]

for item in indice:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# ============ INTRODUÇÃO ============
doc.add_heading('1️⃣ INTRODUÇÃO', level=1)

intro_text = """Este guia foi criado especialmente para você que é iniciante em TI. Vamos instalar e configurar o APOGEU passo a passo, como se estivéssemos conversando pessoalmente.

O que você vai aprender:
• Instalar programas necessários (Node.js, Git, MySQL)
• Baixar o código do projeto
• Configurar o banco de dados
• Iniciar a aplicação
• Acessar a plataforma no navegador
• Resolver problemas comuns e menos frequentes

Tempo estimado: 1-2 horas (primeira vez)

Importante: Não se preocupe se não entender alguns termos - explicaremos tudo!"""

doc.add_paragraph(intro_text)

doc.add_heading('⚠️ Erros que Você Pode Encontrar', level=2)

errors_intro = """Durante a instalação, você pode encontrar erros. Não se preocupe! Criamos uma seção especial com:
✅ Erros comuns (que a maioria encontra)
✅ Erros menos frequentes (mais raros)
✅ Soluções passo a passo para cada um

Procure pelo erro que recebeu e siga a solução."""

doc.add_paragraph(errors_intro)

doc.add_page_break()

# ============ PRÉ-REQUISITOS ============
doc.add_heading('2️⃣ PRÉ-REQUISITOS DO SISTEMA', level=1)

doc.add_heading('Verificar Versão do Windows', level=2)

windows_check = """1. Clique no botão Windows (canto inferior esquerdo)
2. Digite: winver e pressione Enter
3. Uma janela aparecerá mostrando sua versão
4. Procure por "Windows 10" e confirme que diz "Pro"

Se não for Pro: Não é problema! Você pode usar a versão Home também."""

doc.add_paragraph(windows_check)

doc.add_heading('Desativar Antivírus Temporariamente', level=2)

antivirus_text = """⚠️ IMPORTANTE: Durante a instalação, o antivírus pode bloquear alguns arquivos.

Como desativar o Windows Defender:
1. Clique no botão Windows
2. Digite: Segurança do Windows e abra
3. Clique em Proteção contra vírus e ameaças
4. Clique em Gerenciar configurações
5. Desative Proteção em tempo real
6. Clique em Sim quando perguntado

⚠️ REATIVE o antivírus quando terminar a instalação!"""

doc.add_paragraph(antivirus_text)

doc.add_page_break()

# ============ NODE.JS ============
doc.add_heading('3️⃣ INSTALAR NODE.JS', level=1)

doc.add_heading('O que é Node.js?', level=2)
doc.add_paragraph('Node.js é um programa que permite executar código JavaScript no seu computador (não apenas no navegador).')

doc.add_heading('Passo 1: Baixar Node.js', level=2)

nodejs_download = """1. Abra seu navegador (Chrome, Edge, Firefox, etc)
2. Acesse: https://nodejs.org/
3. Você verá dois botões grandes:
   • LTS (versão estável - recomendada) ← CLIQUE AQUI
   • Current (versão nova)
4. Um arquivo .msi começará a baixar
5. Procure o arquivo em C:\\Users\\SeuUsuário\\Downloads"""

doc.add_paragraph(nodejs_download)

doc.add_heading('Passo 2: Instalar Node.js', level=2)

nodejs_install = """1. Abra a pasta Downloads
2. Encontre o arquivo node-v20.x.x-x64.msi
3. Clique duas vezes para iniciar a instalação
4. Uma janela de instalação aparecerá:
   • Clique em Next (Próximo)
   • Leia e aceite os termos: marque a caixa e clique Next
   • Escolha o local (deixe como padrão): clique Next
   • Clique Next novamente
   • ⭐ IMPORTANTE: Marque "Automatically install the necessary tools"
   • Clique Next
   • Clique Install
   • Aguarde 5-10 minutos
   • Clique Finish

5. Uma janela do PowerShell pode aparecer:
   • Pressione Y (Sim) e Enter
   • Aguarde a conclusão"""

doc.add_paragraph(nodejs_install)

doc.add_heading('Passo 3: Verificar Instalação', level=2)

nodejs_verify = """1. Clique no botão Windows
2. Digite: PowerShell e abra
3. Uma janela preta aparecerá
4. Digite: node --version e pressione Enter
5. Você deve ver algo como: v20.10.0

Se não funcionar:
✓ Feche o PowerShell
✓ Reinicie o computador
✓ Tente novamente"""

doc.add_paragraph(nodejs_verify)

doc.add_heading('❌ ERROS COMUNS - NODE.JS', level=2)

# Tabela de erros
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'ERRO'
header_cells[1].text = 'SOLUÇÃO'

# Erro 1
row1 = table.rows[1].cells
row1[0].text = '"node" is not recognized'
row1[1].text = 'Reinicie o computador. O Windows precisa atualizar as variáveis de ambiente.'

# Erro 2
row2 = table.rows[2].cells
row2[0].text = 'Permission denied'
row2[1].text = 'Abra PowerShell como Administrador. Clique direito e escolha "Run as administrator"'

# Erro 3
row3 = table.rows[3].cells
row3[0].text = 'Installation failed'
row3[1].text = 'Desative antivírus, limpe pasta Downloads, tente novamente'

doc.add_page_break()

# ============ GIT ============
doc.add_heading('4️⃣ INSTALAR GIT', level=1)

doc.add_heading('O que é Git?', level=2)
doc.add_paragraph('Git é um programa que permite baixar código de repositórios online (como GitHub).')

doc.add_heading('Passo 1: Baixar Git', level=2)

git_download = """1. Abra seu navegador
2. Acesse: https://git-scm.com/download/win
3. Um arquivo .exe começará a baixar automaticamente
4. Procure por Git-2.x.x-64-bit.exe em Downloads"""

doc.add_paragraph(git_download)

doc.add_heading('Passo 2: Instalar Git', level=2)

git_install = """1. Clique duas vezes no arquivo Git-2.x.x-64-bit.exe
2. Uma janela de instalação aparecerá:
   • Clique Next várias vezes
   • Quando perguntado sobre "Select Components", deixe padrão
   • Clique Next até "Configuring line ending conversions"
   • Deixe: "Checkout Windows-style, commit Unix-style"
   • Clique Next até o final
   • Clique Install
   • Aguarde a instalação
   • Clique Finish"""

doc.add_paragraph(git_install)

doc.add_heading('Passo 3: Verificar Instalação', level=2)

git_verify = """1. Abra PowerShell novamente
2. Digite: git --version e pressione Enter
3. Você deve ver: git version 2.42.0.windows.1"""

doc.add_paragraph(git_verify)

doc.add_heading('❌ ERROS COMUNS - GIT', level=2)

table2 = doc.add_table(rows=3, cols=2)
table2.style = 'Light Grid Accent 1'

header_cells2 = table2.rows[0].cells
header_cells2[0].text = 'ERRO'
header_cells2[1].text = 'SOLUÇÃO'

row1_2 = table2.rows[1].cells
row1_2[0].text = '"git" is not recognized'
row1_2[1].text = 'Reinicie o computador e tente novamente'

row2_2 = table2.rows[2].cells
row2_2[0].text = 'Installation failed'
row2_2[1].text = 'Desative antivírus e tente novamente'

doc.add_page_break()

# ============ MYSQL ============
doc.add_heading('5️⃣ INSTALAR MYSQL', level=1)

doc.add_heading('O que é MySQL?', level=2)
doc.add_paragraph('MySQL é um banco de dados - um lugar para armazenar informações da sua aplicação.')

doc.add_heading('Passo 1: Baixar MySQL', level=2)

mysql_download = """1. Abra seu navegador
2. Acesse: https://dev.mysql.com/downloads/mysql/
3. Escolha a versão 8.0 (a mais estável)
4. Clique em Download
5. Na próxima página: "No thanks, just start my download"
6. Um arquivo .msi começará a baixar"""

doc.add_paragraph(mysql_download)

doc.add_heading('Passo 2: Instalar MySQL', level=2)

mysql_install = """1. Clique duas vezes no arquivo mysql-installer-community-8.x.x.x.msi
2. Uma janela de instalação aparecerá:
   • Clique Next
   • Escolha "Server only" (apenas o servidor)
   • Clique Next
   • Clique Execute para instalar
   • Clique Next quando terminar
   • Clique Next novamente
   • Escolha "Standalone MySQL Server / Classic MySQL Server"
   • Clique Next
   • Deixe Port: 3306 (padrão)
   • Clique Next
   • Escolha "MySQL Server as a Windows Service"
   • Clique Next
   • Deixe nome como "MySQL80"
   • Clique Next
   • Deixe "Standard System Account" selecionado
   • Clique Next

3. ⭐ IMPORTANTE - MySQL Root Password:
   • Digite uma senha que você não vai esquecer
   • Exemplo: Senha123!
   • Confirme a senha no campo abaixo
   • Clique Next
   • Clique Execute
   • Clique Finish"""

doc.add_paragraph(mysql_install)

doc.add_heading('Passo 3: Verificar Instalação', level=2)

mysql_verify = """1. Abra PowerShell
2. Digite: mysql --version e pressione Enter
3. Você deve ver: mysql Ver 8.0.35 for Win64"""

doc.add_paragraph(mysql_verify)

doc.add_heading('❌ ERROS COMUNS - MYSQL', level=2)

table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Light Grid Accent 1'

header_cells3 = table3.rows[0].cells
header_cells3[0].text = 'ERRO'
header_cells3[1].text = 'SOLUÇÃO'

row1_3 = table3.rows[1].cells
row1_3[0].text = '"mysql" is not recognized'
row1_3[1].text = 'Reinicie o computador. Verifique se MySQL está rodando: Services (Win+R, services.msc)'

row2_3 = table3.rows[2].cells
row2_3[0].text = 'Access denied for user root'
row2_3[1].text = 'Você digitou a senha errada. Reinstale MySQL e escolha uma senha simples'

row3_3 = table3.rows[3].cells
row3_3[0].text = 'Port 3306 already in use'
row3_3[1].text = 'Outro MySQL está rodando. Desinstale e reinstale, ou mude a porta'

doc.add_page_break()

# ============ PNPM ============
doc.add_heading('6️⃣ INSTALAR PNPM', level=1)

doc.add_heading('O que é pnpm?', level=2)
doc.add_paragraph('pnpm é um gerenciador de pacotes - um programa que instala bibliotecas que o APOGEU precisa.')

doc.add_heading('Passo 1: Instalar pnpm', level=2)

pnpm_install = """1. Abra PowerShell
2. Digite este comando e pressione Enter:

npm install -g pnpm

3. Aguarde a instalação (1-2 minutos)
4. Você verá mensagens de progresso"""

doc.add_paragraph(pnpm_install)

doc.add_heading('Passo 2: Verificar Instalação', level=2)

pnpm_verify = """1. Feche e abra PowerShell novamente
2. Digite: pnpm --version e pressione Enter
3. Você deve ver: 8.15.0 (ou versão similar)"""

doc.add_paragraph(pnpm_verify)

doc.add_heading('❌ ERROS COMUNS - PNPM', level=2)

table4 = doc.add_table(rows=3, cols=2)
table4.style = 'Light Grid Accent 1'

header_cells4 = table4.rows[0].cells
header_cells4[0].text = 'ERRO'
header_cells4[1].text = 'SOLUÇÃO'

row1_4 = table4.rows[1].cells
row1_4[0].text = '"pnpm" is not recognized'
row1_4[1].text = 'Feche e abra PowerShell novamente. Se persistir, reinicie o computador'

row2_4 = table4.rows[2].cells
row2_4[0].text = 'Permission denied'
row2_4[1].text = 'Abra PowerShell como Administrador'

doc.add_page_break()

# ============ BAIXAR CÓDIGO ============
doc.add_heading('7️⃣ BAIXAR CÓDIGO DO APOGEU', level=1)

doc.add_heading('Passo 1: Criar Pasta para o Projeto', level=2)

folder_create = """1. Abra o Explorador de Arquivos (ícone de pasta na barra de tarefas)
2. Clique em "Este Computador"
3. Abra a unidade C:
4. Clique com botão direito em espaço vazio
5. Clique em "Nova Pasta"
6. Digite: Projetos e pressione Enter
7. Abra a pasta Projetos que você criou"""

doc.add_paragraph(folder_create)

doc.add_heading('Passo 2: Baixar Repositório', level=2)

repo_download = """1. Abra PowerShell
2. Digite e pressione Enter:

cd C:\\Projetos

3. Agora digite este comando e pressione Enter:

git clone https://github.com/keday49c/flower-bloom-network.git

4. Você verá mensagens:
   Cloning into 'flower-bloom-network'...
   remote: Counting objects: 100%

5. Aguarde até ver: done.

6. Agora entre na pasta:

cd flower-bloom-network"""

doc.add_paragraph(repo_download)

doc.add_heading('Passo 3: Verificar Download', level=2)

verify_download = """1. Abra Explorador de Arquivos
2. Navegue até: C:\\Projetos\\flower-bloom-network
3. Você deve ver pastas como: client, server, drizzle, etc."""

doc.add_paragraph(verify_download)

doc.add_heading('❌ ERROS COMUNS - DOWNLOAD', level=2)

table5 = doc.add_table(rows=4, cols=2)
table5.style = 'Light Grid Accent 1'

header_cells5 = table5.rows[0].cells
header_cells5[0].text = 'ERRO'
header_cells5[1].text = 'SOLUÇÃO'

row1_5 = table5.rows[1].cells
row1_5[0].text = '"git" is not recognized'
row1_5[1].text = 'Git não foi instalado corretamente. Reinstale Git'

row2_5 = table5.rows[2].cells
row2_5[0].text = 'fatal: unable to access repository'
row2_5[1].text = 'Problema de internet. Verifique conexão e tente novamente'

row3_5 = table5.rows[3].cells
row3_5[0].text = 'Permission denied'
row3_5[1].text = 'Abra PowerShell como Administrador'

doc.add_page_break()

# ============ INSTALAR DEPENDÊNCIAS ============
doc.add_heading('8️⃣ INSTALAR DEPENDÊNCIAS', level=1)

doc.add_heading('O que são Dependências?', level=2)
doc.add_paragraph('Dependências são bibliotecas de código que o APOGEU precisa para funcionar.')

doc.add_heading('Passo 1: Instalar Dependências', level=2)

deps_install = """1. Abra PowerShell
2. Certifique-se de estar em: C:\\Projetos\\flower-bloom-network
3. Digite este comando e pressione Enter:

pnpm install

4. Você verá:
   Resolving dependencies...
   Fetching packages...

5. ⚠️ IMPORTANTE: Este processo pode levar 10-20 minutos
6. NÃO FECHE o PowerShell enquanto estiver instalando
7. Quando terminar, você verá: Done in X.XXs"""

doc.add_paragraph(deps_install)

doc.add_heading('❌ ERROS COMUNS - DEPENDÊNCIAS', level=2)

table6 = doc.add_table(rows=4, cols=2)
table6.style = 'Light Grid Accent 1'

header_cells6 = table6.rows[0].cells
header_cells6[0].text = 'ERRO'
header_cells6[1].text = 'SOLUÇÃO'

row1_6 = table6.rows[1].cells
row1_6[0].text = 'permission denied'
row1_6[1].text = 'Abra PowerShell como Administrador (clique direito → Run as administrator)'

row2_6 = table6.rows[2].cells
row2_6[0].text = 'ENOENT: no such file'
row2_6[1].text = 'Você não está na pasta correta. Digite: cd C:\\Projetos\\flower-bloom-network'

row3_6 = table6.rows[3].cells
row3_6[0].text = 'ERR! code ERESOLVE'
row3_6[1].text = 'Problema com versão do Node. Desinstale Node e reinstale versão LTS'

doc.add_page_break()

# ============ BANCO DE DADOS ============
doc.add_heading('9️⃣ CONFIGURAR BANCO DE DADOS', level=1)

doc.add_heading('Passo 1: Criar Banco de Dados', level=2)

db_create = """1. Abra PowerShell
2. Digite este comando e pressione Enter:

mysql -u root -p

3. Você será pedido a senha:
   • Digite a senha do MySQL que você criou (ex: Senha123!)
   • Pressione Enter

4. Você verá: mysql>

5. Digite este comando e pressione Enter:

CREATE DATABASE apogeu;

6. Você verá: Query OK, 1 row affected

7. Digite este comando e pressione Enter:

CREATE USER 'apogeu_user'@'localhost' IDENTIFIED BY 'senha_apogeu_123';

8. Você verá: Query OK, 0 rows affected

9. Digite este comando e pressione Enter:

GRANT ALL PRIVILEGES ON apogeu.* TO 'apogeu_user'@'localhost';

10. Você verá: Query OK, 0 rows affected

11. Digite este comando e pressione Enter:

FLUSH PRIVILEGES;

12. Você verá: Query OK, 0 rows affected

13. Digite: EXIT; e pressione Enter para sair"""

doc.add_paragraph(db_create)

doc.add_heading('Passo 2: Verificar Banco', level=2)

db_verify = """1. Abra PowerShell
2. Digite:

mysql -u apogeu_user -p apogeu

3. Digite a senha: senha_apogeu_123
4. Se você ver mysql>, significa que funcionou!
5. Digite EXIT; para sair"""

doc.add_paragraph(db_verify)

doc.add_heading('❌ ERROS COMUNS - BANCO DE DADOS', level=2)

table7 = doc.add_table(rows=5, cols=2)
table7.style = 'Light Grid Accent 1'

header_cells7 = table7.rows[0].cells
header_cells7[0].text = 'ERRO'
header_cells7[1].text = 'SOLUÇÃO'

row1_7 = table7.rows[1].cells
row1_7[0].text = '"mysql" is not recognized'
row1_7[1].text = 'MySQL não está instalado ou não está no PATH. Reinstale MySQL'

row2_7 = table7.rows[2].cells
row2_7[0].text = 'Access denied for user root'
row2_7[1].text = 'Você digitou a senha errada. Tente a senha que você criou durante instalação'

row3_7 = table7.rows[3].cells
row3_7[0].text = 'ERROR 1045 (28000)'
row3_7[1].text = 'Problema de autenticação. Verifique senha e tente novamente'

row4_7 = table7.rows[4].cells
row4_7[0].text = 'Can\'t connect to MySQL'
row4_7[1].text = 'MySQL não está rodando. Abra Services (Win+R, services.msc) e inicie MySQL80'

doc.add_page_break()

# ============ VARIÁVEIS DE AMBIENTE ============
doc.add_heading('🔟 CONFIGURAR VARIÁVEIS DE AMBIENTE', level=1)

doc.add_heading('O que são Variáveis de Ambiente?', level=2)
doc.add_paragraph('São configurações que a aplicação precisa para funcionar (como senhas, chaves de API, etc).')

doc.add_heading('Passo 1: Criar Arquivo .env', level=2)

env_create = """1. Abra Explorador de Arquivos
2. Navegue até: C:\\Projetos\\flower-bloom-network
3. Procure por um arquivo chamado .env.example
4. Clique com botão direito nele
5. Clique em "Abrir com" → "Bloco de Notas"

6. O arquivo abrirá. Você verá:
   # Database
   DATABASE_URL=mysql://user:password@localhost:3306/apogeu
   ...

7. Clique em Arquivo → Salvar Como

8. Na janela que aparecer:
   • No campo "Nome do arquivo", apague tudo e digite: .env
   • No campo "Tipo", escolha: "Todos os arquivos (*.*)"
   • Clique em Salvar

9. Agora você tem um arquivo .env na pasta"""

doc.add_paragraph(env_create)

doc.add_heading('Passo 2: Editar Arquivo .env', level=2)

env_edit = """1. Clique com botão direito no arquivo .env
2. Escolha "Abrir com" → "Bloco de Notas"

3. Encontre esta linha:
   DATABASE_URL=mysql://user:password@localhost:3306/apogeu

4. Substitua por:
   DATABASE_URL=mysql://apogeu_user:senha_apogeu_123@localhost:3306/apogeu

5. Deixe as outras linhas como estão por enquanto

6. Clique em Arquivo → Salvar
7. Feche o Bloco de Notas"""

doc.add_paragraph(env_edit)

doc.add_heading('❌ ERROS COMUNS - VARIÁVEIS', level=2)

table8 = doc.add_table(rows=3, cols=2)
table8.style = 'Light Grid Accent 1'

header_cells8 = table8.rows[0].cells
header_cells8[0].text = 'ERRO'
header_cells8[1].text = 'SOLUÇÃO'

row1_8 = table8.rows[1].cells
row1_8[0].text = 'Arquivo não salva como .env'
row1_8[1].text = 'Certifique-se de escolher "Todos os arquivos (*.*)" no tipo de arquivo'

row2_8 = table8.rows[2].cells
row2_8[0].text = 'Arquivo aparece como .env.txt'
row2_8[1].text = 'Renomeie: clique direito → Renomear, remova .txt'

doc.add_page_break()

# ============ MIGRAÇÕES ============
doc.add_heading('1️⃣1️⃣ EXECUTAR MIGRAÇÕES', level=1)

doc.add_heading('O que é Migração?', level=2)
doc.add_paragraph('É um processo que cria as tabelas e estruturas necessárias no banco de dados.')

doc.add_heading('Passo 1: Executar Migrações', level=2)

migrations = """1. Abra PowerShell
2. Certifique-se de estar em: C:\\Projetos\\flower-bloom-network
3. Digite este comando e pressione Enter:

pnpm db:push

4. Você verá:
   Generating migrations...
   Applying migrations...

5. Aguarde até ver: ✓ Done

6. Se tudo correu bem, suas tabelas foram criadas!"""

doc.add_paragraph(migrations)

doc.add_heading('❌ ERROS COMUNS - MIGRAÇÕES', level=2)

table9 = doc.add_table(rows=3, cols=2)
table9.style = 'Light Grid Accent 1'

header_cells9 = table9.rows[0].cells
header_cells9[0].text = 'ERRO'
header_cells9[1].text = 'SOLUÇÃO'

row1_9 = table9.rows[1].cells
row1_9[0].text = 'Database connection failed'
row1_9[1].text = 'Verifique se MySQL está rodando e se a senha no .env está correta'

row2_9 = table9.rows[2].cells
row2_9[0].text = 'Access denied'
row2_9[1].text = 'Verifique a senha no arquivo .env'

doc.add_page_break()

# ============ BACKEND ============
doc.add_heading('1️⃣2️⃣ INICIAR BACKEND', level=1)

doc.add_heading('O que é Backend?', level=2)
doc.add_paragraph('É a parte "invisível" da aplicação que processa dados.')

doc.add_heading('Passo 1: Iniciar Backend', level=2)

backend = """1. Abra PowerShell
2. Certifique-se de estar em: C:\\Projetos\\flower-bloom-network
3. Digite este comando e pressione Enter:

pnpm dev

4. Você verá:
   Starting dev server...
   [OAuth] Initialized with baseURL: https://api.manus.im
   Server running on http://localhost:3000/

5. ⭐ NÃO FECHE ESTE POWERSHELL! Deixe rodando

6. Se tudo funcionou, você verá a mensagem final em verde"""

doc.add_paragraph(backend)

doc.add_heading('❌ ERROS COMUNS - BACKEND', level=2)

table10 = doc.add_table(rows=4, cols=2)
table10.style = 'Light Grid Accent 1'

header_cells10 = table10.rows[0].cells
header_cells10[0].text = 'ERRO'
header_cells10[1].text = 'SOLUÇÃO'

row1_10 = table10.rows[1].cells
row1_10[0].text = 'Cannot find module'
row1_10[1].text = 'Dependências não foram instaladas. Execute: pnpm install'

row2_10 = table10.rows[2].cells
row2_10[0].text = 'Port 3000 already in use'
row2_10[1].text = 'Outro programa usa porta 3000. Feche ou mude a porta'

row3_10 = table10.rows[3].cells
row3_10[0].text = 'Database connection failed'
row3_10[1].text = 'MySQL não está rodando. Verifique e inicie MySQL'

doc.add_page_break()

# ============ FRONTEND ============
doc.add_heading('1️⃣3️⃣ INICIAR FRONTEND', level=1)

doc.add_heading('O que é Frontend?', level=2)
doc.add_paragraph('É a parte visual da aplicação que você vê no navegador.')

doc.add_heading('Passo 1: Abrir Novo PowerShell', level=2)

frontend_new = """1. ⭐ NÃO FECHE o PowerShell anterior! Deixe rodando
2. Clique no botão Windows
3. Digite: PowerShell e abra um NOVO PowerShell"""

doc.add_paragraph(frontend_new)

doc.add_heading('Passo 2: Navegar para Pasta do Cliente', level=2)

frontend_nav = """1. No novo PowerShell, digite:

cd C:\\Projetos\\flower-bloom-network\\client

2. Pressione Enter"""

doc.add_paragraph(frontend_nav)

doc.add_heading('Passo 3: Instalar Dependências do Cliente', level=2)

frontend_deps = """1. Digite este comando e pressione Enter:

pnpm install

2. Aguarde a instalação (5-10 minutos)"""

doc.add_paragraph(frontend_deps)

doc.add_heading('Passo 4: Iniciar Frontend', level=2)

frontend_start = """1. Digite este comando e pressione Enter:

pnpm dev

2. Você verá:
   VITE v7.1.7 ready in 234 ms
   ➜ Local: http://localhost:5173/

3. ⭐ NÃO FECHE ESTE POWERSHELL! Deixe rodando"""

doc.add_paragraph(frontend_start)

doc.add_page_break()

# ============ ACESSAR APLICAÇÃO ============
doc.add_heading('1️⃣4️⃣ ACESSAR A APLICAÇÃO', level=1)

doc.add_heading('Passo 1: Abrir no Navegador', level=2)

browser = """1. Abra seu navegador (Chrome, Edge, Firefox, etc)
2. Na barra de endereço, digite:

http://localhost:5173

3. Pressione Enter

4. Você deve ver:
   ✅ Logo do APOGEU
   ✅ Página de login
   ✅ Interface profissional com cores azul, roxo e vinho"""

doc.add_paragraph(browser)

doc.add_heading('Passo 2: Fazer Login', level=2)

login = """1. Clique em "Entrar" ou "Login"
2. Você será redirecionado para a página de autenticação
3. Faça login com suas credenciais
4. Você pode criar uma conta se não tiver"""

doc.add_paragraph(login)

doc.add_heading('Passo 3: Explorar a Aplicação', level=2)

explore = """Parabéns! Você está dentro do APOGEU! Agora você pode:

✅ Ver o Dashboard com gráficos e métricas
✅ Criar Campanhas de marketing
✅ Configurar Credenciais de APIs
✅ Visualizar Analytics em tempo real
✅ Gerenciar CRM com leads
✅ Analisar Concorrentes
✅ Ver Planos de Assinatura"""

doc.add_paragraph(explore)

doc.add_heading('❌ ERROS COMUNS - NAVEGADOR', level=2)

table11 = doc.add_table(rows=4, cols=2)
table11.style = 'Light Grid Accent 1'

header_cells11 = table11.rows[0].cells
header_cells11[0].text = 'ERRO'
header_cells11[1].text = 'SOLUÇÃO'

row1_11 = table11.rows[1].cells
row1_11[0].text = 'Página em branco'
row1_11[1].text = 'Pressione F5 para recarregar. Verifique console (F12) por erros'

row2_11 = table11.rows[2].cells
row2_11[0].text = 'Conexão recusada'
row2_11[1].text = 'Backend não está rodando. Verifique PowerShell do backend'

row3_11 = table11.rows[3].cells
row3_11[0].text = 'Página carrega lentamente'
row3_11[1].text = 'Normal na primeira vez. Aguarde. Se persistir, reinicie ambos PowerShells'

doc.add_page_break()

# ============ TESTES ============
doc.add_heading('1️⃣5️⃣ EXECUTAR TESTES', level=1)

doc.add_heading('O que são Testes?', level=2)
doc.add_paragraph('São verificações automáticas para garantir que tudo está funcionando corretamente.')

doc.add_heading('Passo 1: Executar Testes de Segurança', level=2)

tests = """1. Abra um NOVO PowerShell
2. Navegue até: C:\\Projetos\\flower-bloom-network
3. Digite este comando e pressione Enter:

pnpm test

4. Você verá:
   Test Files  1 passed (1)
        Tests  32 passed (32)

5. Todos os 32 testes de segurança devem passar!"""

doc.add_paragraph(tests)

doc.add_heading('Passo 2: Verificar TypeScript', level=2)

typescript = """1. No mesmo PowerShell, digite:

pnpm check

2. Se não houver erros, você verá uma mensagem de sucesso"""

doc.add_paragraph(typescript)

doc.add_page_break()

# ============ APIS EXTERNAS ============
doc.add_heading('1️⃣6️⃣ CONFIGURAR APIS EXTERNAS', level=1)

doc.add_heading('Importante: Isto é OPCIONAL', level=2)
doc.add_paragraph('Se você quiser usar todas as funcionalidades do APOGEU, precisa configurar algumas APIs. Isto é opcional para começar.')

doc.add_heading('Google Ads API', level=2)

google_ads = """Como obter a chave:
1. Abra seu navegador
2. Acesse: https://ads.google.com/
3. Faça login com sua conta Google
4. Clique em Ferramentas → API do Google Ads
5. Siga as instruções para gerar uma chave
6. Copie a chave

Como adicionar ao .env:
1. Abra o arquivo .env com Bloco de Notas
2. Encontre: GOOGLE_ADS_API_KEY=your_google_ads_api_key
3. Substitua por: GOOGLE_ADS_API_KEY=sua_chave_aqui
4. Salve o arquivo"""

doc.add_paragraph(google_ads)

doc.add_heading('OpenAI API (Para IA)', level=2)

openai = """Como obter a chave:
1. Abra seu navegador
2. Acesse: https://platform.openai.com/
3. Crie uma conta ou faça login
4. Clique em API keys
5. Clique em Create new secret key
6. Copie a chave (aparece apenas uma vez!)

Como adicionar ao .env:
1. Abra o arquivo .env com Bloco de Notas
2. Encontre: OPENAI_API_KEY=your_openai_api_key
3. Substitua por: OPENAI_API_KEY=sk-proj-sua_chave_aqui
4. Salve o arquivo"""

doc.add_paragraph(openai)

doc.add_heading('ElevenLabs API (Para Voz)', level=2)

elevenlabs = """Como obter a chave:
1. Abra seu navegador
2. Acesse: https://elevenlabs.io/
3. Crie uma conta
4. Clique em Profile → API Key
5. Copie a chave

Como adicionar ao .env:
1. Abra o arquivo .env com Bloco de Notas
2. Encontre: ELEVENLABS_API_KEY=your_elevenlabs_api_key
3. Substitua por: ELEVENLABS_API_KEY=sua_chave_aqui
4. Salve o arquivo"""

doc.add_paragraph(elevenlabs)

doc.add_heading('Twilio (Para WhatsApp)', level=2)

twilio = """Como obter as credenciais:
1. Abra seu navegador
2. Acesse: https://www.twilio.com/
3. Crie uma conta
4. Clique em Console
5. Você verá:
   • Account SID
   • Auth Token
6. Copie ambos

Como adicionar ao .env:
1. Abra o arquivo .env com Bloco de Notas
2. Encontre:
   TWILIO_ACCOUNT_SID=your_twilio_sid
   TWILIO_AUTH_TOKEN=your_twilio_token
3. Substitua pelos seus valores
4. Salve o arquivo"""

doc.add_paragraph(twilio)

doc.add_heading('Reiniciar Aplicação Após Adicionar Chaves', level=2)

restart_api = """1. Feche o PowerShell que está rodando pnpm dev (Ctrl + C)
2. Digite novamente: pnpm dev
3. Atualize o navegador (F5)"""

doc.add_paragraph(restart_api)

doc.add_page_break()

# ============ BUILD ============
doc.add_heading('1️⃣7️⃣ BUILD PARA PRODUÇÃO', level=1)

doc.add_heading('O que é Build?', level=2)
doc.add_paragraph('É o processo de preparar a aplicação para ser usada em um servidor real.')

doc.add_heading('Passo 1: Fazer Build', level=2)

build = """1. Abra um NOVO PowerShell
2. Navegue até: C:\\Projetos\\flower-bloom-network
3. Digite este comando e pressione Enter:

pnpm build

4. Você verá mensagens de progresso
5. Aguarde até ver: ✓ built in X.XXs

6. Você verá duas pastas criadas:
   • dist (código frontend compilado)
   • dist-electron (código Electron compilado)"""

doc.add_paragraph(build)

doc.add_page_break()

# ============ DESKTOP ============
doc.add_heading('1️⃣8️⃣ CRIAR APLICATIVO DESKTOP', level=1)

doc.add_heading('O que é Electron?', level=2)
doc.add_paragraph('É uma tecnologia que permite usar a aplicação web como um programa desktop (como um .exe).')

doc.add_heading('Passo 1: Instalar Dependências Electron', level=2)

electron_deps = """1. Abra PowerShell
2. Navegue até: C:\\Projetos\\flower-bloom-network
3. Digite este comando e pressione Enter:

pnpm add -D electron electron-builder

4. Aguarde a instalação"""

doc.add_paragraph(electron_deps)

doc.add_heading('Passo 2: Criar Build Desktop', level=2)

electron_build = """1. No mesmo PowerShell, digite:

pnpm electron-build-win

2. Você verá:
   Building for Windows...

3. Aguarde até ver: ✓ Done

4. Os arquivos serão criados em:
   C:\\Projetos\\flower-bloom-network\\dist\\electron

5. Você terá:
   • APOGEU Setup 1.0.0.exe (instalador - recomendado)
   • APOGEU 1.0.0.exe (versão portável)"""

doc.add_paragraph(electron_build)

doc.add_heading('Passo 3: Instalar o Aplicativo', level=2)

electron_install = """1. Abra Explorador de Arquivos
2. Navegue até: C:\\Projetos\\flower-bloom-network\\dist\\electron
3. Clique duas vezes em APOGEU Setup 1.0.0.exe
4. Siga o assistente:
   • Clique Next
   • Escolha pasta de instalação (deixe padrão)
   • Clique Install
   • Clique Finish

5. O APOGEU será instalado como um programa normal!
6. Você pode abrir pelo Menu Iniciar"""

doc.add_paragraph(electron_install)

doc.add_page_break()

# ============ SOLUCIONAR PROBLEMAS ============
doc.add_heading('1️⃣9️⃣ SOLUCIONAR PROBLEMAS', level=1)

doc.add_heading('Problema: "Cannot find module"', level=2)

problem1 = """Significado: Dependências não foram instaladas corretamente

Solução:
1. Abra PowerShell
2. Navegue até a pasta do projeto
3. Digite:

Remove-Item -Recurse -Force node_modules
pnpm install

4. Aguarde a reinstalação"""

doc.add_paragraph(problem1)

doc.add_heading('Problema: "Port 3000 already in use"', level=2)

problem2 = """Significado: Outro programa está usando a porta 3000

Solução:
1. Abra PowerShell como Administrador
2. Digite:

netstat -ano | findstr :3000

3. Você verá um número (PID)
4. Digite:

taskkill /PID [numero] /F

5. Substitua [numero] pelo número que você viu"""

doc.add_paragraph(problem2)

doc.add_heading('Problema: "Database connection failed"', level=2)

problem3 = """Significado: Não consegue conectar ao MySQL

Solução:
1. Verifique se MySQL está rodando
2. Abra PowerShell
3. Digite:

mysql -u root -p

4. Digite sua senha
5. Se funcionar, MySQL está ok
6. Verifique se a senha no .env está correta"""

doc.add_paragraph(problem3)

doc.add_heading('Problema: "Cannot find git"', level=2)

problem4 = """Significado: Git não foi instalado corretamente

Solução:
1. Reinicie o computador
2. Abra um novo PowerShell
3. Tente novamente: git --version"""

doc.add_paragraph(problem4)

doc.add_heading('Problema: Página em branco no navegador', level=2)

problem5 = """Solução:
1. Pressione F12 para abrir console do navegador
2. Procure por mensagens de erro em vermelho
3. Copie a mensagem de erro
4. Procure na internet ou abra issue no GitHub"""

doc.add_paragraph(problem5)

doc.add_heading('❌ ERROS MENOS FREQUENTES', level=2)

doc.add_heading('Erro: "EACCES: permission denied"', level=3)

error_rare1 = """Causa: Problema de permissões no Windows
Solução:
1. Abra PowerShell como Administrador
2. Tente novamente
3. Se persistir, reinicie o computador"""

doc.add_paragraph(error_rare1)

doc.add_heading('Erro: "ENOMEM: out of memory"', level=3)

error_rare2 = """Causa: Seu computador não tem memória suficiente
Solução:
1. Feche outros programas
2. Reinicie o computador
3. Tente novamente
4. Se persistir, você pode precisar de mais RAM"""

doc.add_paragraph(error_rare2)

doc.add_heading('Erro: "ENOENT: no such file or directory"', level=3)

error_rare3 = """Causa: Arquivo ou pasta não encontrada
Solução:
1. Verifique se está na pasta correta
2. Digite: cd C:\\Projetos\\flower-bloom-network
3. Tente novamente"""

doc.add_paragraph(error_rare3)

doc.add_page_break()

# ============ PRÓXIMAS ETAPAS ============
doc.add_heading('2️⃣0️⃣ PRÓXIMAS ETAPAS', level=1)

doc.add_heading('Parabéns! Você tem o APOGEU rodando!', level=2)

next_steps = """Agora você pode:

1. Explorar a Plataforma:
   • Crie campanhas de marketing
   • Configure credenciais de APIs
   • Veja gráficos e análises
   • Teste o bot do WhatsApp
   • Analise concorrentes

2. Configurar Mais APIs:
   • Google Ads para gerenciar anúncios
   • Meta Ads para Facebook/Instagram
   • TikTok Ads para TikTok
   • OpenAI para IA

3. Fazer Deploy em Produção:
   • Contrate um servidor (Heroku, AWS, DigitalOcean)
   • Siga o guia de deploy da plataforma
   • Configure variáveis de ambiente
   • Faça o deploy do código

4. Aprender Mais:
   • Leia a documentação em README.md
   • Explore o código em client/src
   • Veja os serviços em server/services"""

doc.add_paragraph(next_steps)

doc.add_heading('Precisa de Ajuda?', level=2)

help_text = """Se tiver problemas:

1. Abra uma issue no GitHub:
   • Acesse: https://github.com/keday49c/flower-bloom-network/issues
   • Clique em New Issue
   • Descreva seu problema
   • Inclua mensagens de erro

2. Procure na internet:
   • Copie a mensagem de erro
   • Cole no Google
   • Geralmente alguém já teve o mesmo problema

3. Peça ajuda:
   • Comunidades de desenvolvimento
   • Discord de Node.js
   • Stack Overflow"""

doc.add_paragraph(help_text)

doc.add_page_break()

# ============ CHECKLIST FINAL ============
doc.add_heading('✅ CHECKLIST FINAL', level=1)

checklist_items = [
    'Node.js instalado',
    'Git instalado',
    'MySQL instalado',
    'pnpm instalado',
    'Código baixado',
    'Dependências instaladas',
    'Banco de dados criado',
    'Arquivo .env configurado',
    'Migrações executadas',
    'Backend rodando (pnpm dev)',
    'Frontend rodando (pnpm dev)',
    'Aplicação acessível em http://localhost:5173',
    'Login funcionando',
    'Testes passando',
    'Build criado',
    'Aplicativo desktop instalado',
]

for item in checklist_items:
    p = doc.add_paragraph('☐ ' + item, style='List Bullet')

doc.add_page_break()

# ============ CONCLUSÃO ============
doc.add_heading('🎉 CONCLUSÃO', level=1)

conclusion = """Você completou a instalação do APOGEU! 

Você agora tem uma plataforma profissional de marketing digital rodando no seu computador!

Próximo passo: Explore a plataforma, crie campanhas e divirta-se! 🚀

APOGEU - Seu Pico de Sucesso em Marketing Digital!"""

final_para = doc.add_paragraph(conclusion)
final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Salvar documento
doc.save('/home/ubuntu/apogeu/GUIA_COMPLETO_APOGEU.docx')
print('✅ Documento Word criado com sucesso!')
print('📄 Arquivo: GUIA_COMPLETO_APOGEU.docx')

